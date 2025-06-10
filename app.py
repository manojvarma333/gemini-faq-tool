import os
import json
import requests
import time
from flask import Flask, render_template, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
import google.generativeai as genai
from dotenv import load_dotenv
import PyPDF2
from docx import Document
import markdown
from bs4 import BeautifulSoup
import re

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)

# Configuration
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'md'}

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize Gemini API with timeout configuration
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            if len(pdf_reader.pages) == 0:
                raise Exception("PDF file has no pages")
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as page_error:
                    print(f"Warning: Could not extract text from page {page_num + 1}: {page_error}")
                    continue
            
            if not text.strip():
                raise Exception("No text content found in PDF")
                
            return text
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        
        if len(doc.paragraphs) == 0:
            raise Exception("DOCX file has no paragraphs")
        
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        if not text.strip():
            raise Exception("No text content found in DOCX")
            
        return text
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")

def extract_text_from_file(file_path, file_extension):
    """Extract text from various file types"""
    try:
        if file_extension == 'pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension == 'docx':
            return extract_text_from_docx(file_path)
        elif file_extension == 'md':
            # Try different encodings for markdown files
            encodings = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        return markdown.markdown(content)
                except UnicodeDecodeError:
                    continue
            raise Exception("Could not decode markdown file with any supported encoding")
        else:  # txt
            # Try different encodings for text files
            encodings = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise Exception("Could not decode text file with any supported encoding")
    except Exception as e:
        raise Exception(f"Failed to extract text from {file_extension} file: {str(e)}")

def scrape_website_content(url):
    """Scrape content from a website"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text content
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return text
    except Exception as e:
        return f"Error scraping website: {str(e)}"

def get_available_models():
    """Get list of available Gemini models"""
    try:
        models = genai.list_models()
        return [model.name for model in models if 'generateContent' in model.supported_generation_methods]
    except Exception as e:
        print(f"Error getting models: {e}")
        return []

def generate_faq_with_gemini(content, num_questions=10, style="professional", max_retries=3):
    """Generate FAQ using Google Gemini API with retry logic"""
    
    # Prepare the prompt based on style
    style_prompts = {
        "professional": "Generate a professional FAQ section with clear, concise questions and detailed answers.",
        "casual": "Generate a friendly, conversational FAQ section with approachable language.",
        "technical": "Generate a technical FAQ section with detailed explanations and code examples where relevant.",
        "simple": "Generate a simple, easy-to-understand FAQ section with basic explanations."
    }
    
    prompt = f"""
    Based on the following content, create {num_questions} frequently asked questions and their answers.
    
    Style: {style_prompts.get(style, style_prompts["professional"])}
    
    Content:
    {content[:4000]}  # Limit content to avoid token limits
    
    Please format the response as a JSON array with the following structure:
    [
        {{
            "question": "Question text here",
            "answer": "Detailed answer here"
        }}
    ]
    
    Make sure the questions are relevant to the content and cover the most important topics.
    Only return the JSON array, no additional text or explanations.
    """
    
    # Try different model names in order of preference
    model_names = [
        'gemini-1.5-pro',
        'gemini-1.5-flash',
        'gemini-pro',
        'gemini-1.0-pro'
    ]
    
    for attempt in range(max_retries):
        for model_name in model_names:
            try:
                print(f"Attempting with model: {model_name}")
                
                # Initialize Gemini model with timeout
                model = genai.GenerativeModel(model_name)
                
                # Generate response using Gemini with timeout
                response = model.generate_content(
                    prompt,
                    generation_config=genai.types.GenerationConfig(
                        temperature=0.7,
                        max_output_tokens=2000,
                    )
                )
                
                response_text = response.text.strip()
                
                # Try to extract JSON from the response
                try:
                    # Look for JSON array in the response
                    json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
                    if json_match:
                        faq_data = json.loads(json_match.group())
                        return faq_data
                    else:
                        # If no JSON found, try to parse the entire response
                        faq_data = json.loads(response_text)
                        return faq_data
                except json.JSONDecodeError:
                    # Fallback: create a simple structure
                    return [{"question": "Generated FAQ", "answer": response_text}]
                    
            except Exception as e:
                error_msg = str(e)
                print(f"Model {model_name} failed: {error_msg}")
                
                # Check if it's a model not found error
                if "not found" in error_msg.lower() or "404" in error_msg:
                    continue  # Try next model
                
                # Check if it's a timeout or connection error
                if "timeout" in error_msg.lower() or "connection" in error_msg.lower() or "503" in error_msg:
                    if attempt < max_retries - 1:
                        print(f"Attempt {attempt + 1} failed: {error_msg}")
                        print(f"Retrying in {2 ** attempt} seconds...")
                        time.sleep(2 ** attempt)  # Exponential backoff
                        break  # Try next attempt with same model
                    else:
                        return [{
                            "question": "Network Error", 
                            "answer": f"Unable to connect to Gemini API after {max_retries} attempts. Please check your internet connection and try again. Error: {error_msg}"
                        }]
                else:
                    # For other errors, try next model
                    continue
    
    return [{"question": "Error", "answer": "Failed to generate FAQ. No available models found or all attempts failed."}]

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_faq', methods=['POST'])
def generate_faq():
    try:
        data = request.get_json()
        content = data.get('content', '')
        num_questions = int(data.get('num_questions', 10))
        style = data.get('style', 'professional')
        
        if not content.strip():
            return jsonify({"error": "No content provided"}), 400
        
        # Generate FAQ using Gemini
        faq_data = generate_faq_with_gemini(content, num_questions, style)
        
        return jsonify({
            "success": True,
            "faq": faq_data,
            "count": len(faq_data)
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/upload_file', methods=['POST'])
def upload_file():
    try:
        print("File upload request received")
        
        if 'file' not in request.files:
            print("No file in request.files")
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        print(f"File received: {file.filename}")
        
        if file.filename == '':
            print("Empty filename")
            return jsonify({"error": "No file selected"}), 400
        
        if not file:
            print("File object is None")
            return jsonify({"error": "Invalid file object"}), 400
        
        if not allowed_file(file.filename):
            print(f"Invalid file type: {file.filename}")
            return jsonify({"error": f"Invalid file type. Supported types: {', '.join(ALLOWED_EXTENSIONS)}"}), 400
        
        # Create secure filename
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        print(f"Saving file to: {file_path}")
        
        # Save the file
        file.save(file_path)
        
        # Check if file was saved successfully
        if not os.path.exists(file_path):
            print("File was not saved successfully")
            return jsonify({"error": "Failed to save file"}), 500
        
        # Extract text from file
        file_extension = filename.rsplit('.', 1)[1].lower()
        print(f"Extracting text from {file_extension} file")
        
        try:
            content = extract_text_from_file(file_path, file_extension)
            print(f"Extracted content length: {len(content)} characters")
            
            if not content or content.strip() == "":
                return jsonify({"error": "No text content found in file"}), 400
                
        except Exception as extract_error:
            print(f"Error extracting text: {extract_error}")
            # Clean up the uploaded file
            if os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({"error": f"Failed to extract text from file: {str(extract_error)}"}), 500
        
        # Clean up the uploaded file
        try:
            os.remove(file_path)
            print("Temporary file cleaned up")
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up temporary file: {cleanup_error}")
        
        return jsonify({
            "success": True,
            "content": content,
            "filename": filename,
            "content_length": len(content)
        })
        
    except Exception as e:
        print(f"Unexpected error in file upload: {e}")
        return jsonify({"error": f"File upload failed: {str(e)}"}), 500

@app.route('/scrape_website', methods=['POST'])
def scrape_website():
    try:
        data = request.get_json()
        url = data.get('url', '')
        
        if not url:
            return jsonify({"error": "No URL provided"}), 400
        
        # Scrape website content
        content = scrape_website_content(url)
        
        return jsonify({
            "success": True,
            "content": content,
            "url": url
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/export_faq', methods=['POST'])
def export_faq():
    try:
        data = request.get_json()
        faq_data = data.get('faq', [])
        format_type = data.get('format', 'html')
        
        if format_type == 'html':
            html_content = """
            <!DOCTYPE html>
            <html>
            <head>
                <title>Generated FAQ</title>
                <style>
                    body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }
                    .faq-item { margin-bottom: 30px; }
                    .question { font-weight: bold; color: #333; margin-bottom: 10px; }
                    .answer { color: #666; line-height: 1.6; }
                </style>
            </head>
            <body>
                <h1>Frequently Asked Questions</h1>
            """
            
            for item in faq_data:
                html_content += f"""
                <div class="faq-item">
                    <div class="question">Q: {item['question']}</div>
                    <div class="answer">A: {item['answer']}</div>
                </div>
                """
            
            html_content += "</body></html>"
            
            return jsonify({
                "success": True,
                "content": html_content,
                "format": "html"
            })
        
        elif format_type == 'json':
            return jsonify({
                "success": True,
                "content": json.dumps(faq_data, indent=2),
                "format": "json"
            })
        
        else:
            return jsonify({"error": "Unsupported format"}), 400
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000) 